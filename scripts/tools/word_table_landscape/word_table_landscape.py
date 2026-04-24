#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word 表格横排工具
功能：将 .docx 文档页面方向改为横向A4，并对所有表格进行格式美化
输出：原文件名_横排.docx（同目录）

格式规范（对标 Excel格式整理工具）：
  - 表头行：灰色背景 #D9D9D9 + 微软雅黑 11pt 加粗 + 居中 + 自动换行
  - 数据行：微软雅黑 11pt + 左对齐 + 顶对齐 + 自动换行
  - 列宽：按列数平均分配横版A4可用宽度
  - 表格边框：外框加粗，内部细线
"""

import sys
import os
import copy
from pathlib import Path

from docx import Document
from docx.shared import Mm, Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import lxml.etree as etree

# ──────────────────────────────────────────────
# 常量
# ──────────────────────────────────────────────
HEADER_BG_COLOR = "D9D9D9"   # 表头背景色（与 Excel工具一致）
FONT_NAME       = "微软雅黑"
FONT_SIZE_PT    = 11

# 横向A4可用宽度（mm）：297 - 左边距25 - 右边距25 = 247mm
PAGE_W_MM       = 297
PAGE_H_MM       = 210
MARGIN_MM       = 25
USABLE_W_MM     = PAGE_W_MM - MARGIN_MM * 2   # 247mm


# ──────────────────────────────────────────────
# 辅助函数：边框 XML
# ──────────────────────────────────────────────
def make_border_element(tag, val="single", sz="6", space="0", color="000000"):
    """构造一个 w:XXX 边框子元素"""
    el = OxmlElement(tag)
    el.set(qn("w:val"),   val)
    el.set(qn("w:sz"),    sz)
    el.set(qn("w:space"), space)
    el.set(qn("w:color"), color)
    return el


def set_table_borders(table):
    """
    给表格设置边框：
      外框：粗线（sz=12 即 1.5pt）
      内部：细线（sz=6  即 0.75pt）
    """
    tbl  = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    # 移除旧 tblBorders（如有）
    old = tblPr.find(qn("w:tblBorders"))
    if old is not None:
        tblPr.remove(old)

    tblBorders = OxmlElement("w:tblBorders")
    # 外框（top / left / bottom / right）粗
    for side in ("top", "left", "bottom", "right"):
        tblBorders.append(make_border_element(f"w:{side}", sz="12"))
    # 内部（insideH / insideV）细
    for side in ("insideH", "insideV"):
        tblBorders.append(make_border_element(f"w:{side}", sz="6"))
    tblPr.append(tblBorders)


def set_cell_borders(cell):
    """给单元格设置细边框（内部线，保持一致性）"""
    tc    = cell._tc
    tcPr  = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)

    old = tcPr.find(qn("w:tcBorders"))
    if old is not None:
        tcPr.remove(old)

    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        tcBorders.append(make_border_element(f"w:{side}", sz="6"))
    tcPr.append(tcBorders)


# ──────────────────────────────────────────────
# 辅助函数：单元格背景色
# ──────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    """设置单元格背景色，hex_color 不含 #，如 'D9D9D9'"""
    tc   = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)

    old = tcPr.find(qn("w:shd"))
    if old is not None:
        tcPr.remove(old)

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


# ──────────────────────────────────────────────
# 辅助函数：段落格式
# ──────────────────────────────────────────────
def format_paragraph(para, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    """设置段落字体、对齐；处理已有 run 并确保有一个 run"""
    para.alignment = align
    # 段落级别 spacing
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(2)

    # 若段落没有 run，新建一个占位 run
    if not para.runs:
        run = para.add_run()
    else:
        run = None

    for r in para.runs:
        r.font.name       = FONT_NAME
        r.font.size       = Pt(FONT_SIZE_PT)
        r.font.bold       = bold
        r.font.color.rgb  = RGBColor(0x00, 0x00, 0x00)
        # 东亚字体
        rPr = r._r.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), FONT_NAME)
        rFonts.set(qn("w:ascii"),    FONT_NAME)
        rFonts.set(qn("w:hAnsi"),    FONT_NAME)

    if run is not None:
        run.font.name      = FONT_NAME
        run.font.size      = Pt(FONT_SIZE_PT)
        run.font.bold      = bold
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)


# ──────────────────────────────────────────────
# 辅助函数：设置单元格垂直对齐
# ──────────────────────────────────────────────
def set_cell_valign(cell, valign: str = "top"):
    """valign: 'top' | 'center' | 'bottom'"""
    tc   = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    old = tcPr.find(qn("w:vAlign"))
    if old is not None:
        tcPr.remove(old)
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), valign)
    tcPr.append(vAlign)


# ──────────────────────────────────────────────
# 辅助函数：设置单元格自动换行
# ──────────────────────────────────────────────
def set_cell_wrap(cell):
    """确保单元格开启自动换行（Word 默认开启，但显式设置更稳）"""
    tc   = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    # 移除 noWrap（若有）
    old = tcPr.find(qn("w:noWrap"))
    if old is not None:
        tcPr.remove(old)


# ──────────────────────────────────────────────
# 核心函数：美化单张表格
# ──────────────────────────────────────────────
def beautify_table(table, col_width_mm: float):
    """
    对一张表格进行格式美化：
      - 第一行视为表头（灰底+加粗+居中）
      - 其余行为数据行（无底色+左对齐）
      - 列宽均等分配
      - 设置边框
    """
    col_width_twips = int(Mm(col_width_mm))  # Mm() 返回 EMU，需转换

    # python-docx 的 Mm() 返回 EMU (English Metric Units)
    # 列宽用 Twips：1 inch = 1440 twips = 914400 EMU → 1 twip = 635 EMU
    col_width_emu = Mm(col_width_mm)  # EMU

    # 设置表格整体对齐
    tbl  = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    # 表格宽度：自动（按列宽撑满）
    old_w = tblPr.find(qn("w:tblW"))
    if old_w is not None:
        tblPr.remove(old_w)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"),    "0")
    tblW.set(qn("w:type"), "auto")
    tblPr.append(tblW)

    # 边框
    set_table_borders(table)

    num_cols = len(table.columns)

    for row_idx, row in enumerate(table.rows):
        is_header = (row_idx == 0)

        for col_idx, cell in enumerate(row.cells):
            # 列宽
            tc   = cell._tc
            tcPr = tc.find(qn("w:tcPr"))
            if tcPr is None:
                tcPr = OxmlElement("w:tcPr")
                tc.insert(0, tcPr)
            old_cw = tcPr.find(qn("w:tcW"))
            if old_cw is not None:
                tcPr.remove(old_cw)
            tcW = OxmlElement("w:tcW")
            # 使用 dxa 单位（twips），1 EMU = 1/635 twip
            twips_val = int(col_width_emu / 635)
            tcW.set(qn("w:w"),    str(twips_val))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)

            # 背景色
            if is_header:
                set_cell_bg(cell, HEADER_BG_COLOR)
            # 垂直对齐
            set_cell_valign(cell, "center" if is_header else "top")
            # 自动换行
            set_cell_wrap(cell)
            # 单元格边框
            set_cell_borders(cell)

            # 段落格式
            for para in cell.paragraphs:
                if is_header:
                    format_paragraph(para, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
                else:
                    format_paragraph(para, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT)


# ──────────────────────────────────────────────
# 核心函数：设置文档为横向A4
# ──────────────────────────────────────────────
def set_landscape_a4(doc: Document):
    """将文档所有节设置为横向A4，并统一边距"""
    for section in doc.sections:
        section.orientation  = WD_ORIENT.LANDSCAPE
        section.page_width   = Mm(PAGE_W_MM)
        section.page_height  = Mm(PAGE_H_MM)
        section.left_margin  = Mm(MARGIN_MM)
        section.right_margin = Mm(MARGIN_MM)
        section.top_margin   = Mm(20)
        section.bottom_margin = Mm(20)


# ──────────────────────────────────────────────
# 主处理函数
# ──────────────────────────────────────────────
def process_file(filepath: str):
    src = Path(filepath).expanduser().resolve()
    if not src.exists():
        print(f"[错误] 文件不存在：{src}", file=sys.stderr)
        return False

    if src.suffix.lower() == ".doc":
        print(f"[错误] {src.name} 是 .doc 格式，请先转换为 .docx")
        return False

    dst = src.parent / f"{src.stem}_横排{src.suffix}"

    # 避免覆盖已存在文件
    counter = 1
    original_dst = dst
    while dst.exists():
        dst = original_dst.parent / f"{src.stem}_横排_{counter}{src.suffix}"
        counter += 1

    try:
        doc = Document(str(src))
    except Exception as e:
        print(f"[错误] 无法打开文档：{e}", file=sys.stderr)
        return False

    # 1. 改为横向A4
    set_landscape_a4(doc)

    # 2. 美化所有表格（不过滤列数，用户自己选文件）
    tables = doc.tables
    if not tables:
        print(f"[提示] 文档中未发现表格：{src.name}")
    else:
        for i, table in enumerate(tables):
            num_cols = len(table.columns)
            if num_cols == 0:
                continue
            col_w = USABLE_W_MM / num_cols
            beautify_table(table, col_w)
            print(f"  ✓ 表格 {i+1}：{num_cols} 列，每列 {col_w:.1f}mm")

    # 3. 保存
    try:
        doc.save(str(dst))
        print(f"[完成] 已保存：{dst.name}")
        return True
    except Exception as e:
        print(f"[错误] 保存失败：{e}", file=sys.stderr)
        return False


# ──────────────────────────────────────────────
# 入口
# ──────────────────────────────────────────────
if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法：python3 word_table_landscape.py <file1.docx> [file2.docx ...]")
        sys.exit(1)

    success = True
    for fp in sys.argv[1:]:
        ok = process_file(fp)
        if not ok:
            success = False

    sys.exit(0 if success else 1)
