#!/usr/bin/env python3
"""
MD转Word工具 - 优化版
将Markdown文件转换为Word文档，保留加粗格式，清除斜体格式，减号作为纯文本
"""

import sys
import re
import os
import shutil
import tempfile
import subprocess
import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── 通用 Helper：向段落添加带格式的 runs ─────────────────────
def _add_formatted_runs(paragraph, text, bold=False, font_name='微软雅黑',
                         font_size=Pt(10), color=RGBColor(0x00, 0x00, 0x00)):
    """向段落添加带格式的 runs，自动处理 **bold** 标记。"""
    if not text:
        return

    if '**' in text:
        parts = re.split(r'(\*\*[^*]+\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.font.bold = True
            else:
                if part:
                    run = paragraph.add_run(part)
                    if bold:
                        run.font.bold = True
    else:
        run = paragraph.add_run(text)
        if bold:
            run.font.bold = True

    for r in paragraph.runs:
        r.font.name = font_name
        if font_name == '微软雅黑':
            r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        r.font.size = font_size
        r.font.color.rgb = color


def clean_markdown_formatting(text):
    """清除Markdown格式符号，保留加粗标记用于后续处理，清除斜体。"""
    bold_placeholders = []
    def save_bold(match):
        bold_placeholders.append(match.group(1))
        return f"\x00BOLD{len(bold_placeholders)-1}\x00"

    text = re.sub(r'\*\*([^*]+)\*\*', save_bold, text)
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    text = re.sub(r'_([^_]+)_', r'\1', text)

    for i, content in enumerate(bold_placeholders):
        text = text.replace(f"\x00BOLD{i}\x00", f"**{content}**")

    text = re.sub(r'`([^`]+)`', r'\1', text)
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    text = re.sub(r'!\[([^\]]*)\]\([^\)]+\)', r'【\1】', text)
    text = text.replace('[', '【').replace(']', '】')
    return text


def apply_formatting_to_run(run, text):
    """应用格式到run，处理加粗标记（保留向后兼容）。"""
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run_bold = run._parent.add_run(part[2:-2])
            run_bold.font.name = run.font.name
            run_bold._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run_bold.font.size = run.font.size
            run_bold.font.bold = True
            run_bold.font.color.rgb = run.font.color.rgb
        else:
            if part:
                run_normal = run._parent.add_run(part)
                run_normal.font.name = run.font.name
                run_normal._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                run_normal.font.size = run.font.size
                run_normal.font.color.rgb = run.font.color.rgb
    run._element.getparent().remove(run._element)


def set_cell_shading(cell, fill_color):
    """设置单元格背景色。"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def is_separator_row(cells):
    """判断是否为Markdown表格分隔行（支持对齐标记 :）。"""
    if not cells:
        return False
    separator_pattern = re.compile(r'^[\s\-:]+$')
    return all(separator_pattern.match(cell) for cell in cells)


def parse_table_line(line):
    """解析表格行，处理转义的管道符。"""
    temp_placeholder = '\x00PIPE\x00'
    line = line.replace('\\|', temp_placeholder)
    cells = [cell.strip() for cell in line.split('|')[1:-1]]
    cells = [cell.replace(temp_placeholder, '|') for cell in cells]
    return cells


def create_beautiful_table(doc, table_data):
    """创建美观的表格（支持列数不一致）。"""
    if len(table_data) < 1:
        return

    max_cols = max(len(row) for row in table_data)
    if max_cols == 0:
        return

    normalized_data = []
    for row in table_data:
        if len(row) < max_cols:
            row = row + [''] * (max_cols - len(row))
        elif len(row) > max_cols:
            row = row[:max_cols]
        normalized_data.append(row)

    cleaned_data = []
    for row in normalized_data:
        cleaned_row = [clean_markdown_formatting(cell) for cell in row]
        cleaned_data.append(cleaned_row)

    table = doc.add_table(rows=len(cleaned_data), cols=max_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False

    for i, row_data in enumerate(cleaned_data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = ''
            _add_formatted_runs(cell.paragraphs[0], cell_text)

            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                if i == 0:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    header_row = table.rows[0]
    for cell in header_row.cells:
        set_cell_shading(cell, 'D3D3D3')

    for row in table.rows:
        row.height = Pt(30)

    doc.add_paragraph()


def process_md_to_docx(md_path):
    """将Markdown转换为Word文档（优化版）。"""
    md_path = Path(md_path)
    if not md_path.exists():
        print(f"❌ 文件不存在: {md_path}")
        return False

    print(f"📄 正在转换: {md_path.name}")

    content = md_path.read_text(encoding='utf-8')
    lines = content.split('\n')

    doc = Document()

    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.font.size = Pt(10)
    style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    in_code_block = False
    code_block_content = []
    in_table = False
    table_data = []

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped_line = line.lstrip()

        if stripped_line.startswith('```'):
            if in_code_block:
                if code_block_content:
                    code_para = doc.add_paragraph()
                    code_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    code_para.paragraph_format.left_indent = Inches(0)
                    code_para.paragraph_format.first_line_indent = Inches(0)
                    for code_line in code_block_content:
                        clean_code = clean_markdown_formatting(code_line)
                        run = code_para.add_run(clean_code + '\n')
                        run.font.name = 'Consolas'
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                    code_block_content = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_block_content.append(line)
            i += 1
            continue

        if stripped_line.startswith('|') and not in_table:
            in_table = True
            table_data = []

        if in_table:
            if stripped_line.startswith('|'):
                cells = parse_table_line(stripped_line)
                if cells and not is_separator_row(cells):
                    table_data.append(cells)
                i += 1
            elif stripped_line == '':
                if i + 1 < len(lines) and lines[i + 1].lstrip().startswith('|'):
                    i += 1
                    continue
                else:
                    if len(table_data) >= 1:
                        create_beautiful_table(doc, table_data)
                    in_table = False
                    table_data = []
            else:
                if len(table_data) >= 1:
                    create_beautiful_table(doc, table_data)
                in_table = False
                table_data = []
        else:
            if not line:
                i += 1
                continue

            if stripped_line == '---' or stripped_line == '***' or stripped_line == '___':
                i += 1
                continue

            if line.startswith('# ') and not line.startswith('## '):
                text = clean_markdown_formatting(line[2:])
                if text:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _add_formatted_runs(p, text, bold=True)
                    doc.add_paragraph()
            elif line.startswith('## '):
                text = clean_markdown_formatting(line[3:])
                if text:
                    p = doc.add_paragraph()
                    _add_formatted_runs(p, text, bold=True)
            elif line.startswith('### '):
                text = clean_markdown_formatting(line[4:])
                if text:
                    p = doc.add_paragraph()
                    _add_formatted_runs(p, text, bold=True)
            elif line.startswith('#### '):
                text = clean_markdown_formatting(line[5:])
                if text:
                    p = doc.add_paragraph()
                    _add_formatted_runs(p, text, bold=True)
            elif line.startswith('- ') or line.startswith('* '):
                text = clean_markdown_formatting(line)
                if text:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0)
                    p.paragraph_format.first_line_indent = Inches(0)
                    _add_formatted_runs(p, text)
            elif re.match(r'^\d+\.\s+', line):
                text = clean_markdown_formatting(line)
                if text:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0)
                    p.paragraph_format.first_line_indent = Inches(0)
                    _add_formatted_runs(p, text)
            elif line.startswith('> '):
                text = clean_markdown_formatting(line[2:])
                if text:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0)
                    p.paragraph_format.first_line_indent = Inches(0)
                    _add_formatted_runs(p, text)
            else:
                text = clean_markdown_formatting(line)
                if text:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0)
                    p.paragraph_format.first_line_indent = Inches(0)
                    _add_formatted_runs(p, text)

            i += 1

    if in_table and len(table_data) >= 1:
        create_beautiful_table(doc, table_data)

    output_path = md_path.with_suffix('.docx')

    counter = 1
    original_path = output_path
    while output_path.exists():
        stem = original_path.stem
        if '_' in stem:
            base_parts = stem.rsplit('_', 1)
            if base_parts[1].isdigit():
                stem = base_parts[0]
        output_path = original_path.parent / f"{stem}_{counter}.docx"
        counter += 1

    current_time = datetime.datetime.now(datetime.timezone.utc)
    doc.core_properties.title = ""
    doc.core_properties.author = os.environ.get('USER', 'User')
    doc.core_properties.comments = ""
    doc.core_properties.keywords = ""
    doc.core_properties.last_modified_by = os.environ.get('USER', 'User')
    doc.core_properties.created = current_time
    doc.core_properties.modified = current_time

    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
            tmp_path = tmp_file.name

        doc.save(tmp_path)

        try:
            result = subprocess.run(['cp', '-X', tmp_path, str(output_path)],
                                    capture_output=True, text=True, timeout=5)
            if result.returncode != 0:
                with open(tmp_path, 'rb') as src, open(str(output_path), 'wb') as dst:
                    dst.write(src.read())

            for attr in ['com.apple.quarantine', 'com.apple.provenance',
                         'com.apple.macl', 'com.apple.metadata:kMDItemWhereFroms']:
                try:
                    subprocess.run(['xattr', '-d', attr, str(output_path)],
                                   capture_output=True, check=False, timeout=2)
                except Exception:
                    pass

            os.chmod(str(output_path), 0o644)

        except Exception as e:
            print(f"警告：文件处理时出错: {e}")
            os.replace(tmp_path, str(output_path))

    finally:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.unlink(tmp_path)
            except Exception:
                pass

    if output_path != original_path:
        print(f"✅ 转换完成: {output_path}（原文件 {original_path.name} 已存在，自动重命名）")
    else:
        print(f"✅ 转换完成: {output_path}")
    return True


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("用法: python md2docx_plain.py <markdown文件>")
        sys.exit(1)

    success_count = 0
    for arg in sys.argv[1:]:
        try:
            if process_md_to_docx(arg):
                success_count += 1
        except Exception as e:
            print(f"❌ 处理失败 [{arg}]: {e}", file=sys.stderr)

    print(f"\n总计: 成功 {success_count}/{len(sys.argv)-1}")
