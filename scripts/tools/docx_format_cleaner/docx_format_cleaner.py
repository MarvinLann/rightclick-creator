#!/usr/bin/env python3
"""
DOCX 格式整理工具 - 通用版（无硬编码路径）
流程：docx → md → 清洗中间MD → md → docx → 后处理清洗
效果：清除原 Word 中的杂乱格式、修订痕迹，统一为干净排版。
"""

import sys
import os
import re
import shutil
import tempfile
import subprocess
from pathlib import Path

# 脚本自定位（支持任意安装路径）
INSTALL_DIR = Path(__file__).parent
DOCX2MD_SCRIPT = INSTALL_DIR / "docx2md_converter.py"
MD2DOCX_SCRIPT = INSTALL_DIR / "md2docx_plain.py"

# ---- 正则：emoji ----
EMOJI_PATTERN = re.compile(
    "["
    "\U0001F1E0-\U0001F1FF"
    "\U0001F300-\U0001F5FF"
    "\U0001F600-\U0001F64F"
    "\U0001F680-\U0001F6FF"
    "\U0001F700-\U0001F77F"
    "\U0001F780-\U0001F7FF"
    "\U0001F800-\U0001F8FF"
    "\U0001F900-\U0001F9FF"
    "\U0001FA00-\U0001FA6F"
    "\U0001FA70-\U0001FAFF"
    "\U00002702-\U000027B0"
    "\U000024C2-\U00002B55"
    "\u2600-\u26FF"
    "\u2700-\u27BF"
    "\uFE00-\uFE0F"
    "\u200D"
    "\u20E3"
    "]+", flags=re.UNICODE)


def _find_soffice() -> str:
    """动态探测 LibreOffice/soffice 路径"""
    candidates = [
        "/opt/homebrew/bin/soffice",
        "/usr/local/bin/soffice",
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
    ]
    for p in candidates:
        if Path(p).exists():
            return p
    found = shutil.which("soffice")
    if found:
        return found
    raise RuntimeError(
        "未找到 LibreOffice。处理 .doc 文件需要 LibreOffice：\n"
        "  brew install --cask libreoffice\n"
        "或从 https://www.libreoffice.org 下载安装。"
    )


def clean_intermediate_md(md_path: str):
    r"""
    清洗中间 MD：
    1. 去反斜杠 —— pandoc 输出会转义 md 符号（如 \*、\[、\-），
       同时会把原文中的单反斜杠 \ 转义为双反斜杠 \\
       策略：去掉所有 pandoc 转义反斜杠，但保留 \n \t \r 等编程转义序列字面量
    2. 去 emoji —— 避免被 md2docx_plain 当正文保留
    """
    content = Path(md_path).read_text(encoding='utf-8')
    # 去掉 pandoc 转义反斜杠，保留 \n \t \r
    content = re.sub(
        r'\\(.)',
        lambda m: m.group(0) if m.group(1) in 'ntr' else m.group(1),
        content,
    )
    content = EMOJI_PATTERN.sub('', content)
    Path(md_path).write_text(content, encoding='utf-8')


def post_clean_docx(docx_path: str):
    """后处理：清洗最终 docx 中残留的 md 排版符号"""
    from docx import Document

    doc = Document(docx_path)

    def clean_text(text: str) -> str:
        # 反斜杠已在 clean_intermediate_md 阶段统一处理，
        # 此处只清理残留的 markdown 格式符号
        text = EMOJI_PATTERN.sub('', text)
        text = re.sub(r'\*\*([^*]*)\*\*', r'\1', text)
        text = re.sub(r'\*([^*]*)\*', r'\1', text)
        text = re.sub(r'^#+\s*', '', text)
        text = text.replace('`', '')
        text = re.sub(r'^>\s*', '', text)
        text = re.sub(r'^-\s+', '', text)
        text = re.sub(r'^\*\s+', '', text)
        text = re.sub(r'^\|\s*', '', text)
        text = re.sub(r'\s*\|$', '', text)
        text = re.sub(r'\s*\|\s*', ' ', text)
        if re.match(r'^[-]{3,}$', text.strip()):
            text = ''
        return text

    for para in doc.paragraphs:
        for run in para.runs:
            new_text = clean_text(run.text)
            if new_text != run.text:
                run.text = new_text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        new_text = clean_text(run.text)
                        if new_text != run.text:
                            run.text = new_text

    doc.save(docx_path)


def doc_to_docx(doc_path: str, output_dir: str) -> str:
    """用 LibreOffice 将 .doc 转为 .docx，返回生成的 docx 路径"""
    soffice = _find_soffice()
    result = subprocess.run(
        [
            soffice,
            "--headless",
            "--convert-to", "docx",
            "--outdir", output_dir,
            doc_path,
        ],
        capture_output=True, text=True, timeout=60,
    )
    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice .doc → .docx 转换失败: {result.stderr}")

    docx_name = Path(doc_path).stem + ".docx"
    docx_path = Path(output_dir) / docx_name
    if not docx_path.exists():
        candidates = list(Path(output_dir).glob("*.docx"))
        if candidates:
            docx_path = candidates[0]
        else:
            raise RuntimeError("LibreOffice 未生成 .docx 文件")

    return str(docx_path)


def format_clean(docx_path: str) -> str:
    """
    对 docx/doc 执行：doc→docx（如需）→ md → 清洗中间MD → md → docx → 后处理清洗
    返回最终输出的 docx 文件路径。
    """
    src = Path(docx_path).expanduser().resolve()
    if not src.exists():
        raise FileNotFoundError(f"文件不存在: {src}")

    if src.suffix.lower() not in (".docx", ".doc"):
        raise ValueError(f"不支持的文件类型: {src.suffix}")

    # 检查工具脚本已安装
    for script in [DOCX2MD_SCRIPT, MD2DOCX_SCRIPT]:
        if not script.exists():
            raise RuntimeError(
                f"工具脚本未找到: {script}\n"
                "请先运行安装：在 AI 助手中触发 docx-rightclick-cleaner Skill 并说「安装右键工具」。"
            )

    with tempfile.TemporaryDirectory() as tmpdir:
        # ---- Step 0: .doc → .docx（如需）----
        if src.suffix.lower() == ".doc":
            tmp_doc_src = Path(tmpdir) / src.name
            shutil.copy2(str(src), str(tmp_doc_src))
            doc_to_docx(str(tmp_doc_src), tmpdir)
            candidates = list(Path(tmpdir).glob("*.docx"))
            if not candidates:
                raise RuntimeError("LibreOffice 转换后未找到 .docx 文件")
            tmp_docx = candidates[0]
        else:
            tmp_docx = Path(tmpdir) / src.name
            shutil.copy2(str(src), str(tmp_docx))

        sys.path.insert(0, str(INSTALL_DIR))
        from docx2md_converter import accept_revisions_and_convert

        md_path = accept_revisions_and_convert(str(tmp_docx))
        md_path = Path(md_path)

        if not md_path.exists():
            raise RuntimeError("docx → md 转换失败，未生成 MD 文件")

        # ---- Step 1.5: 清洗中间MD ----
        clean_intermediate_md(str(md_path))

        # ---- Step 2: md → docx ----
        from md2docx_plain import process_md_to_docx

        work_md = Path(tmpdir) / "work.md"
        shutil.copy2(str(md_path), str(work_md))

        success = process_md_to_docx(str(work_md))
        if not success:
            raise RuntimeError("md → docx 转换失败")

        output_docx = work_md.with_suffix(".docx")
        if not output_docx.exists():
            raise RuntimeError("md → docx 转换失败，未生成 DOCX 文件")

        # ---- Step 2.5: 后处理清洗 ----
        post_clean_docx(str(output_docx))

        # ---- Step 3: 输出新文件（原文件名_整理.docx）----
        new_name = f"{src.stem}_整理.docx"
        dest = src.parent / new_name

        counter = 1
        while dest.exists():
            new_name = f"{src.stem}_整理_{counter}.docx"
            dest = src.parent / new_name
            counter += 1

        result = subprocess.run(
            ["cp", "-X", str(output_docx), str(dest)],
            capture_output=True, text=True, timeout=5,
        )
        if result.returncode != 0:
            shutil.copy2(str(output_docx), str(dest))

        for attr in ["com.apple.quarantine", "com.apple.provenance",
                     "com.apple.metadata:kMDItemWhereFroms"]:
            try:
                subprocess.run(["xattr", "-d", attr, str(dest)],
                               capture_output=True, check=False, timeout=2)
            except Exception:
                pass

        os.chmod(str(dest), 0o644)

    return str(dest)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python3 docx_format_cleaner.py <文件1.docx或.doc> [文件2 ...]")
        sys.exit(1)

    for path in sys.argv[1:]:
        try:
            out = format_clean(path)
            print(f"✅ 格式整理完成: {out}")
        except Exception as e:
            print(f"❌ 处理失败 [{path}]: {e}", file=sys.stderr)
